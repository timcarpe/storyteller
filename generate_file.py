import docx
from docx.shared import Inches

def get_storypage_docx(document, storypage):
    document.add_picture(storypage.image_url,width=Inches(4), height=Inches(4))
    document.add_paragraph(storypage.pagetext, style='Heading 1')
    document.add_paragraph(storypage.question, style='Heading 2')

def generate_document(storypages, filename):
    
    # Create a new Word document
    doc = docx.Document()
    doc.add_heading('Story Teller', 0)

    # Generate the content of the document
    for i, storypage in enumerate(storypages):
        get_storypage_docx(doc, storypage)
        if i != len(storypages)-1:
            doc.add_page_break()

    # Save the document
    filename = filename + ".docx"
    doc.save(filename)
    print("Finished writing: " + filename)